import json
import PyPDF2
import docx
from tqdm import tqdm
import time
import os
import openai
import sys

filename = sys.argv[1] #The Name of the PDf that you are translating. It must have .pdf extension.
model_name = sys.argv[2] #The name of the model you are using. You must look up the approrpiate name.
language = sys.argv[3] #The language you are translating your text to. This will be used in the prompt.

openai.organization = "org-------------"
openai.api_key = "sk-----------------"

pdf = PyPDF2.PdfReader(filename)

start_page = 0
end_page = len(pdf.pages)

doc = docx.Document()
doc.add_heading("This translated is generated using " + str(model_name))

code_crash_page = 0 # In case the code crashes, store the page to restart from.

for pg_num in tqdm(range(start_page, end_page)):
  code_crash_page = pg_num
  pg = pdf.pages[pg_num].extract_text()

  # Send Request for Model Output to inputted model
  success = False
  while not success:
    try:
      response = openai.Completion.create(
        model=model_name,
        prompt=f'''As a professional translator, translate the following text into {language}, while keeping the text structure. Keep the line break smybols in place:\n\n{pg}\n\n ''',
        temperature=0.2,
        max_tokens=2048,
        top_p=1.0,
        frequency_penalty=0.0,
        presence_penalty=0.0
        )
      success = True
    except:
      time.sleep(25)

  # Extract string output from model
  s = json.loads(str(response))['choices'][0]['text']

  # Add string output to document
  doc.add_paragraph(s)
  doc.add_paragraph("")

  if pg_num % 10 == 0:
    doc.save(filename[:-4] + '_' + language + '.docx')
  
doc.save(filename[:-4] + '_' + language + '.docx')
